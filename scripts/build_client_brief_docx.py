from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Billu.Works_Dashboard_V2_Excel_Import_Structure_Brief.docx"


ACCENT = "1F6F64"
ACCENT_2 = "EAF4F1"
TEXT = "1F2933"
MUTED = "5C6B73"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=TEXT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(hdr[index], header, bold=True, color="FFFFFF")
        shade(hdr[index], ACCENT)
        if widths:
            hdr[index].width = widths[index]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
            if widths:
                cells[index].width = widths[index]
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(ACCENT if level <= 2 else TEXT)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Aptos"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(TEXT)
        text = text[len(bold_prefix):]
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(TEXT)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade(cell, ACCENT_2)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Aptos"
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor.from_string(ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    body_run = p2.add_run(body)
    body_run.font.name = "Aptos"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph()


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.color.rgb = RGBColor.from_string(ACCENT)
    return doc


def build():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Billu.Works Dashboard V2")
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Excel Import Structure Brief")
    run.font.name = "Aptos"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Draft for review | May 7, 2026 | Pending the client's actual workbook/sample sheets")
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    add_callout(
        doc,
        "Executive recommendation",
        "Use one master Excel workbook with a clean Transactions sheet containing all months together. Each row should be one transaction, and month-to-month comparison should come from the Date column."
    )

    add_heading(doc, "1. Executive Summary")
    add_body(doc, "The recommended structure is one master Excel workbook, with data organized in clean database-style sheets. For the finance dashboard, the strongest structure is a single Transactions sheet containing all months together.")
    add_body(doc, "This is better than separate custom month layouts because it makes the dashboard easier to compare across months, categories, accounts, vendors, clients, and cash-flow periods.")
    add_body(doc, "The dashboard should not rely on merged parent headers or visual Excel grouping. Instead, hierarchy should be captured in normal columns such as Parent Group, Head / Category, and Subcategory.")

    add_heading(doc, "2. Recommended Workbook Structure")
    add_table(
        doc,
        ["Sheet Name", "Purpose", "Required Now"],
        [
            ["Transactions", "All finance transactions across all months", "Yes"],
            ["Categories / Chart of Accounts", "Optional category list for cleaner mapping and reporting", "Later / optional"],
            ["Opening Balances", "Starting balance by bank/cash account", "Later / optional"],
            ["Notes / Assumptions", "Clarifications for finance/admin review", "Optional"],
        ],
        [Inches(1.8), Inches(4.0), Inches(1.3)],
    )

    add_heading(doc, "3. Recommended Transactions Columns")
    add_table(
        doc,
        ["Column", "Purpose", "Example"],
        [
            ["Date", "Transaction date", "2026-05-01"],
            ["Account", "Bank, cash, wallet, or source account", "HBL Current Account"],
            ["Flow", "Revenue/inflow or expense/outflow", "Revenue"],
            ["Parent Group", "Broad grouping", "Revenue, Expense"],
            ["Head / Category", "Main reporting category", "Sales, Office, Payroll"],
            ["Subcategory", "More detailed child category", "Website Project, Internet"],
            ["Description", "Transaction description/narration", "Client payment for invoice 1021"],
            ["Vendor / Customer", "Counterparty", "ABC Client"],
            ["Amount", "Positive amount", "250000"],
            ["Running Balance", "Account balance after transaction, if available", "1250000"],
            ["Reference / Invoice No", "Invoice, receipt, cheque, or transaction ID", "INV-1021"],
            ["Notes", "Manual explanation", "Advance payment"],
        ],
        [Inches(1.65), Inches(3.35), Inches(2.1)],
    )

    add_callout(
        doc,
        "Minimum required fields",
        "The import can work from Date and Amount, but Account, Flow, Parent Group, Head / Category, Subcategory, Description, Vendor / Customer, and Running Balance will make reporting much stronger."
    )

    add_heading(doc, "4. Parent And Child Headers")
    add_body(doc, "The client mentioned one parent header and multiple child headers. That is common in human-made Excel reports, but it is not the best format for dashboard import.")
    add_body(doc, "Avoid merged or grouped visual headers where the meaning depends on header placement. Use explicit hierarchy columns instead.")
    add_table(
        doc,
        ["Date", "Parent Group", "Head / Category", "Subcategory", "Flow", "Amount"],
        [
            ["2026-05-01", "Revenue", "Sales", "Website Project", "Revenue", "250000"],
            ["2026-05-02", "Revenue", "Retainers", "Monthly Support", "Revenue", "100000"],
            ["2026-05-03", "Expense", "Office", "Rent", "Outflow", "50000"],
            ["2026-05-04", "Expense", "Office", "Internet", "Outflow", "12000"],
        ],
        [Inches(1.0), Inches(1.25), Inches(1.4), Inches(1.45), Inches(1.0), Inches(0.9)],
    )

    add_heading(doc, "5. One Sheet Vs Monthly Sheets")
    add_body(doc, "Best option: one Transactions sheet with all months together. Monthly comparison is automatic from the Date column.")
    add_bullets(doc, [
        "Monthly revenue, expense, and net cash trends become easier to calculate.",
        "Category, subcategory, account, vendor, and customer filters work consistently.",
        "The file is easier to validate and easier for accountant review.",
    ])
    add_body(doc, "Acceptable option: monthly sheets such as Jan 2026, Feb 2026, and Mar 2026. This can work only if every monthly sheet uses the exact same columns and layout.")
    add_body(doc, "If the client prefers monthly sheets, v2 should add a combine-all-monthly-sheets import mode after we review the actual workbook.")

    add_heading(doc, "6. What V2 Can Support Today")
    add_bullets(doc, [
        "CSV import",
        "Excel workbook parsing",
        "Worksheet picking",
        "Header detection",
        "Mapping review",
        "Transaction validation",
        "Finance fields for date, amount, flow, parent/group, category/head, subcategory, description, counterparty, account, and running balance",
    ])
    add_body(doc, "Current v2 is best suited to clean transaction-style data. Once the client's actual workbook is shared, the import flow can be customized around their real column names and office format.")

    add_heading(doc, "7. What We Still Need From The Client")
    add_bullets(doc, [
        "One real or anonymized sample workbook",
        "At least one normal month of data",
        "All columns they expect to use",
        "Any parent/child headers they currently use",
        "Example revenue, expense, transfer, payroll, inventory, and admin rows if those exist",
        "Notes explaining what each column means",
        "A few messy real-world cases, such as missing category, refund, transfer, advance payment, or partial payment",
    ])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "8. Future Platform Direction")
    add_body(doc, "The client's mention of inventory, admin, attendance, and stock is useful. These should become separate platform modules, not extra columns forced into the finance transaction import.")
    add_table(
        doc,
        ["Module", "Suggested Sheet", "What It Tracks"],
        [
            ["Finance", "Transactions", "Revenue, expenses, cash flow, balances"],
            ["Inventory", "Inventory / Stock Ledger", "Items, stock in/out, units, cost, suppliers"],
            ["Attendance", "Attendance", "Employee, date, status, check-in/out, overtime"],
            ["Admin", "Admin Tasks / Operations", "Renewals, documents, licenses, assignments"],
            ["Payroll", "Payroll", "Staff, salary, deductions, payments"],
        ],
        [Inches(1.25), Inches(2.1), Inches(3.75)],
    )

    add_heading(doc, "9. Inventory Module Idea")
    add_body(doc, "Inventory should be handled as a stock ledger rather than only a current-stock list. Every stock movement should be a row.")
    add_table(
        doc,
        ["Column", "Purpose"],
        [
            ["Date", "Stock movement date"],
            ["Item Code", "Unique item/SKU"],
            ["Item Name", "Product/item name"],
            ["Category", "Item group"],
            ["Movement Type", "Purchase, sale, adjustment, return, damage"],
            ["Quantity In", "Stock added"],
            ["Quantity Out", "Stock removed"],
            ["Unit Cost", "Cost per unit"],
            ["Supplier / Customer", "Related party"],
            ["Reference", "Invoice/order number"],
            ["Location", "Warehouse/store/location"],
            ["Notes", "Explanation"],
        ],
        [Inches(2.0), Inches(5.1)],
    )

    add_heading(doc, "10. Attendance Module Idea")
    add_body(doc, "Attendance should also be row-based. This can later support attendance summaries, absences, overtime, payroll preparation, and admin reporting.")
    add_table(
        doc,
        ["Column", "Purpose"],
        [
            ["Date", "Attendance date"],
            ["Employee ID", "Unique employee code"],
            ["Employee Name", "Staff name"],
            ["Department", "Team/department"],
            ["Status", "Present, absent, leave, half-day, remote"],
            ["Check In", "Start time"],
            ["Check Out", "End time"],
            ["Overtime Hours", "Overtime, if any"],
            ["Notes", "Manual explanation"],
        ],
        [Inches(2.0), Inches(5.1)],
    )

    add_heading(doc, "11. Implementation Recommendation")
    add_numbered(doc, [
        "Continue building v2 around the standard transaction model.",
        "Ask the client for one real or anonymized Excel workbook.",
        "Test the workbook in the dashboard.",
        "Identify which columns map cleanly and which need custom logic.",
        "Add import improvements only after seeing the real format.",
        "Keep inventory, attendance, admin, and payroll as future modules with separate data models.",
    ])

    add_heading(doc, "12. Suggested Client Message")
    add_callout(
        doc,
        "Message you can send",
        "For the dashboard, the cleanest option is one master Excel workbook with a Transactions sheet that contains all months together. Each row should be one transaction, and the month should come from the Date column. If you prefer separate monthly sheets, that can also work, but every monthly sheet should use the exact same columns and layout. For parent and child categories, please avoid merged headers or visual grouped headers. Instead, use separate columns like Parent Group, Head / Category, and Subcategory. Once you share a real or anonymized sample workbook, we can test it in the v2 dashboard and customize the import flow around your actual office format."
    )

    add_heading(doc, "13. Bottom Line")
    add_body(doc, "Use one master workbook, one clean Transactions sheet, and explicit hierarchy columns. Keep the office's human-friendly summaries separate from the raw import data. This will make the dashboard easier to build, easier to audit, and easier to extend into future Billu.Works modules like inventory, attendance, payroll, and admin operations.")

    doc.core_properties.author = "Billu.Works"
    doc.core_properties.title = "Billu.Works Dashboard V2 Excel Import Structure Brief"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
